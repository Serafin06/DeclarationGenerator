# services/pdf_generator.py

"""
PDFGenerator - Generuje HTML i PDF z szablonów Jinja2
"""
from jinja2 import Environment, FileSystemLoader
from pathlib import Path
from datetime import datetime
import io
import base64

from src.config.constants import (
    TEMPLATES_PATH, OUTPUT_PATH,
    TEMPLATE_PL_TECH, TEMPLATE_EN_TECH,
    TEMPLATE_PL_BOK, TEMPLATE_EN_BOK
)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from src.models.declaration import Declaration


class PDFGenerator:
    """Generator dokumentów HTML i PDF"""

    def __init__(self, data_loader):
        self.data_loader = data_loader
        self.env = Environment(
            loader=FileSystemLoader(str(TEMPLATES_PATH)),
            autoescape=True
        )
        OUTPUT_PATH.mkdir(exist_ok=True, parents=True)
        self.templates_base_path = TEMPLATES_PATH

    def _get_template_path(self, declaration: Declaration) -> Path:
        """Zwraca odpowiednią ścieżkę szablonu"""
        if declaration.declaration_type == 'tech':
            return TEMPLATE_PL_TECH if declaration.language == 'pl' else TEMPLATE_EN_TECH
        else:
            return TEMPLATE_PL_BOK if declaration.language == 'pl' else TEMPLATE_EN_BOK

    def _prepare_context(self, declaration: Declaration) -> dict:
        """Przygotowuje kontekst dla szablonu"""
        texts = self.data_loader.get_texts(declaration.language)
        context = declaration.to_template_dict()
        context['texts'] = texts

        if declaration.generation_date:
            context['generation_date'] = declaration.generation_date.strftime("%d.%m.%Y")
        else:
            context['generation_date'] = datetime.now().strftime("%d.%m.%Y")

        # Embed obrazy jako base64 (to już masz dobrze)
        logo_path = self.templates_base_path / "logo.jpg"
        podpis_path = self.templates_base_path / "podpis.png"

        if logo_path.exists():
            with open(logo_path, 'rb') as f:
                logo_b64 = base64.b64encode(f.read()).decode()
                context['logo_base64'] = f"data:image/jpeg;base64,{logo_b64}"

        if podpis_path.exists():
            with open(podpis_path, 'rb') as f:
                podpis_b64 = base64.b64encode(f.read()).decode()
                context['podpis_base64'] = f"data:image/png;base64,{podpis_b64}"

        return context

    def generate_html_content(self, declaration: Declaration) -> str:
        """Renderuje szablon do stringa HTML"""
        template_path = self._get_template_path(declaration)
        template = self.env.get_template(template_path.name)
        context = self._prepare_context(declaration)
        return template.render(**context)

    def generate_html(self, declaration: Declaration) -> Path:
        """Zapisuje podgląd HTML i zwraca ścieżkę absolutną"""
        html_content = self.generate_html_content(declaration)
        output_file = (OUTPUT_PATH / "preview_temp.html").resolve()

        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return output_file

    def generate_pdf_bytes(self, declaration: Declaration) -> bytes:
        """Generuje PDF i zwraca jako bajty do zapisu przez użytkownika"""
        html_content = self.generate_html_content(declaration)

        try:
            from weasyprint import HTML
            from pathlib import WindowsPath

            # Konwertuj UNC path na file:// URL dla WeasyPrint
            if isinstance(self.templates_base_path, WindowsPath):
                base_url = self.templates_base_path.as_uri()
            else:
                base_url = str(self.templates_base_path)

            print(f"DEBUG base_url as URI: {base_url}")
            pdf_bytes = HTML(string=html_content, base_url=base_url).write_pdf()
            return pdf_bytes

        except Exception as e:
            raise Exception(f"Błąd generowania PDF: {e}")

    def generate_docx(self, declaration: Declaration, output_path: str):
        """
        Generuje plik DOCX z deklaracji.
        Zapisuje do ścieżki 'output_path' przekazanej z widoku.
        """
        html_content = self.generate_html_content(declaration)
        soup = BeautifulSoup(html_content, 'html.parser')

        # Utwórz dokument Word
        doc = Document()

        # Ustaw marginesy (15mm ~ 0.59 cala)
        section = doc.sections[0]
        section.top_margin = Inches(0.59)
        section.bottom_margin = Inches(0.59)
        section.left_margin = Inches(0.59)
        section.right_margin = Inches(0.59)

        # ===== NAGŁÓWEK (LOGO) =====
        # UWAGA: Aby dodać nagłówek w python-docx, trzeba użyć section.header
        header = section.header
        logo_path = self.templates_base_path / "logo.jpg"
        if logo_path.exists():
            try:
                # Dodaj logo do nagłówka
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                run = header_para.add_run()
                run.add_picture(str(logo_path), width=Inches(1.5))  # ZMIEŃ ROZMIAR TUTAJ
                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception as e:
                print(f"Błąd dodawania logo do nagłówka: {e}")

        # ===== STOPKA =====
        # UWAGA: Dodanie stopki z danymi firmy
        footer = section.footer
        # Linia 1: pogrubiona
        p1 = footer.add_paragraph()
        run1 = p1.add_run("MARPOL Sp. z o.o., Ignatki 40/1, 16-001 Kleosin")
        run1.bold = True  # ZAKOMENTUJ JEŚLI NIE CHCESZ POGRUBIENIA
        run1.font.size = Pt(9)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Linia 2: kontakt
        p2 = footer.add_paragraph("tel. +48 85 74 74 397, fax +48 85 66 31 150, http://www.marpol.pl")
        p2.runs[0].font.size = Pt(7.5)
        p2.runs[0].font.color.rgb = RGBColor(102, 102, 102)  # KOLOR SZARY
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Linia 3: dane rejestrowe
        p3 = footer.add_paragraph(
            "NIP 542-10-04-860, REGON 050297484, KRS 0000857780 | Kapitał zakładowy: 1.167.500,00 PLN")
        p3.runs[0].font.size = Pt(7.5)
        p3.runs[0].font.color.rgb = RGBColor(102, 102, 102)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ===== TREŚĆ DOKUMENTU =====
        body = soup.find('body')
        if body:
            self._process_html_to_docx(doc, body)

        # Zapisz do wybranej przez użytkownika ścieżki
        try:
            doc.save(output_path)
        except Exception as e:
            raise Exception(f"Nie udało się zapisać pliku DOCX: {e}")

    def _process_html_to_docx(self, doc, element):
        """
        Przetwarza elementy HTML do DOCX.
        Obsługuje: nagłówki, akapity, tabele, listy, obrazki.
        """
        for child in element.children:
            if isinstance(child, str):
                text = child.strip()
                if text:
                    doc.add_paragraph(text)
            else:
                tag_name = child.name

                # ===== POMIJANIE ELEMENTÓW NAGŁÓWKA/STOPKI Z HTML =====
                # UWAGA: Elementy z id="header" i id="footer" są już obsłużone w sekcji header/footer
                # więc pomijamy je tutaj, żeby się nie duplikowały
                if child.get('id') in ['header', 'footer']:
                    continue

                # ===== OBRAZKI (LOGO/PODPIS W TREŚCI) =====
                if tag_name == 'img':
                    src = child.get('src', '')
                    try:
                        if src.startswith('data:image'):
                            # Base64 embedded image
                            header, data = src.split(",", 1)
                            img_format = header.split('/')[1].split(';')[0]
                            image_bytes = base64.b64decode(data)
                            stream = io.BytesIO(image_bytes)
                            # ZMIEŃ SZEROKOŚĆ OBRAZKA TUTAJ (np. width=Inches(2.5))
                            doc.add_picture(stream, width=Inches(2.0))
                        else:
                            # Plik lokalny
                            if Path(src).exists():
                                doc.add_picture(src, width=Inches(2.0))
                    except Exception as e:
                        print(f"Błąd wczytywania obrazka do DOCX: {e}")
                    continue

                # ===== NAGŁÓWEK H1 =====
                elif tag_name == 'h1':
                    text = child.get_text().strip()
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.runs[0]
                    run.font.size = Pt(14)
                    # ZAKOMENTUJ PONIŻSZE JEŚLI NIE CHCESZ FORMATOWANIA
                    # run.font.bold = True
                    # run.font.underline = True

                # ===== NAGŁÓWEK H2 =====
                elif tag_name == 'h2':
                    text = child.get_text().strip()
                    p = doc.add_paragraph(text)
                    # ZAKOMENTUJ PONIŻSZE JEŚLI NIE CHCESZ FORMATOWANIA
                    # p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)

                # ===== AKAPITY I DIVY =====
                elif tag_name in ['p', 'div']:
                    if tag_name == 'div':
                        # ZAWSZE wchodź rekurencyjnie w DIV (np. <div class="section">)
                        # aby przetworzyć jego zawartość (tabele, paragrafy, itp.)
                        self._process_html_to_docx(doc, child)
                    else:
                        # Zwykły <p>
                        text = child.get_text().strip()
                        if text:
                            p = doc.add_paragraph(text)

                            # Wyrównanie
                            style = child.get('style', '')
                            if 'text-align: center' in style or 'text-align:center' in style:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif 'text-align: right' in style or 'text-align:right' in style:
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                            # ZAKOMENTUJ PONIŻSZE JEŚLI NIE CHCESZ AUTO-POGRUBIENIA
                            # if child.find('b') or child.find('strong'):
                            #     for run in p.runs:
                            #         run.font.bold = True

                # ===== TABELE =====
                elif tag_name == 'table':
                    self._add_table_to_docx(doc, child)

                # ===== LISTY NUMEROWANE/PUNKTOWANE =====
                elif tag_name == 'ul':
                    for li in child.find_all('li', recursive=False):
                        text = li.get_text().strip()
                        if text:
                            doc.add_paragraph(text, style='List Bullet')

                elif tag_name == 'ol':
                    for li in child.find_all('li', recursive=False):
                        text = li.get_text().strip()
                        if text:
                            doc.add_paragraph(text, style='List Number')

                # ===== LINIE POZIOME I PRZERWY =====
                elif tag_name == 'hr':
                    doc.add_paragraph()  # Pusty wiersz

                elif tag_name == 'br':
                    doc.add_paragraph()

                # ===== INNE ELEMENTY - REKURENCJA =====
                else:
                    self._process_html_to_docx(doc, child)

    def _add_table_to_docx(self, doc, table_element):
        """
        Dodaje tabelę HTML do DOCX.

        UWAGI:
        - word_table.style = 'Table Grid' → tabela Z obramowaniem
        - word_table.style = 'Plain Table 1' → tabela BEZ obramowania (lub zostaw None)
        - Możesz też ręcznie usunąć obramowanie pętlą po komórkach
        """
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Policz maksymalną liczbę kolumn
        cols_count = max(len(r.find_all(['th', 'td'])) for r in rows) if rows else 0
        if cols_count == 0:
            return

        # Utwórz tabelę
        word_table = doc.add_table(rows=len(rows), cols=cols_count)

        # ===== STYL TABELI =====
        # OPCJA 1: Tabela BEZ obramowania (odkomentuj poniższą linię)
        # word_table.style = None

        # OPCJA 2: Tabela Z obramowaniem (odkomentuj poniższą linię)
        word_table.style = 'Table Grid'

        # OPCJA 3: Ręczne usunięcie obramowania (odkomentuj poniższą pętlę)
        # from docx.oxml import OxmlElement
        # from docx.oxml.ns import qn
        # for row in word_table.rows:
        #     for cell in row.cells:
        #         tc = cell._tc
        #         tcPr = tc.get_or_add_tcPr()
        #         tcBorders = OxmlElement('w:tcBorders')
        #         for border_name in ['top', 'left', 'bottom', 'right']:
        #             border = OxmlElement(f'w:{border_name}')
        #             border.set(qn('w:val'), 'none')
        #             tcBorders.append(border)
        #         tcPr.append(tcBorders)

        # Wypełnij komórki
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])

            word_col_idx = 0
            for html_cell in cells:
                if word_col_idx >= cols_count:
                    break

                word_cell = word_table.rows[row_idx].cells[word_col_idx]
                text = html_cell.get_text().strip()

                # Ustaw tekst
                for p in word_cell.paragraphs:
                    p.clear()
                word_cell.paragraphs[0].text = text

                # Rowspan/Colspan
                rowspan = int(html_cell.get('rowspan', 1))
                colspan = int(html_cell.get('colspan', 1))

                # Formatowanie nagłówków
                is_header = (html_cell.name == 'th')

                # Wyrównanie komórki
                style = html_cell.get('style', '')
                align = None
                if 'text-align: center' in style or 'text-align:center' in style:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-align: right' in style or 'text-align:right' in style:
                    align = WD_ALIGN_PARAGRAPH.RIGHT

                # ZAKOMENTUJ PONIŻSZE JEŚLI NIE CHCESZ AUTO-POGRUBIENIA NAGŁÓWKÓW
                # if is_header:
                #     for run in word_cell.paragraphs[0].runs:
                #         run.font.bold = True

                # Scalanie komórek
                if colspan > 1 or rowspan > 1:
                    try:
                        end_row_idx = row_idx + rowspan - 1
                        end_col_idx = word_col_idx + colspan - 1

                        if end_row_idx < len(rows) and end_col_idx < cols_count:
                            end_cell = word_table.rows[end_row_idx].cells[end_col_idx]
                            word_cell.merge(end_cell)
                    except Exception as e:
                        print(f"Błąd scalania komórek: {e}")

                # Zastosuj wyrównanie
                if align:
                    for p in word_cell.paragraphs:
                        p.alignment = align

                word_col_idx += colspan

